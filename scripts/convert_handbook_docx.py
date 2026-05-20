import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    """Sets the background color of a cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets the internal padding (margins) of a cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    """Sets borders for a cell. Colors are hex strings, size in 1/8 pt."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    
    borders = {'top': top, 'bottom': bottom, 'left': left, 'right': right}
    for border_name, border_style in borders.items():
        if border_style:
            color, size, val = border_style
            node = OxmlElement(f'w:{border_name}')
            node.set(qn('w:val'), val)
            node.set(qn('w:sz'), str(size))
            node.set(qn('w:space'), '0')
            node.set(qn('w:color'), color)
            tcBorders.append(node)
        else:
            node = OxmlElement(f'w:{border_name}')
            node.set(qn('w:val'), 'none')
            tcBorders.append(node)
    tcPr.append(tcBorders)

def add_paragraph_with_runs(doc_or_cell, text, style=None, font_name="TH Sarabun New", size_pt=11.5, bold=False, italic=False, color=None, space_after=6, line_spacing=1.15):
    """Adds a paragraph and styles it with run-level formatting, supporting basic markdown bold **."""
    if hasattr(doc_or_cell, 'add_paragraph'):
        p = doc_or_cell.add_paragraph(style=style)
    else:
        p = doc_or_cell.paragraph_format # If it's a paragraph directly
    
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    
    # Split text by bold markers **
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if not part:
            continue
        run = p.add_run()
        run.font.name = font_name
        run.font.size = Pt(size_pt)
        # Apply font name to east asia (Thai) text
        run._r.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), font_name)
        
        if part.startswith('**') and part.endswith('**'):
            run.text = part[2:-2]
            run.bold = True
        else:
            run.text = part
            run.bold = bold
            
        run.italic = italic
        if color:
            run.font.color.rgb = color
            
    return p

def main():
    md_path = "/Applications/XAMPP/xamppfiles/htdocs/rbrupraneet_agri_inno2026/docs/agri_innovation_handbook_3days.md"
    docx_path = "/Applications/XAMPP/xamppfiles/htdocs/rbrupraneet_agri_inno2026/docs/agri_innovation_handbook_3days.docx"
    
    print(f"Reading markdown from: {md_path}")
    if not os.path.exists(md_path):
        print(f"[ERROR] Source Markdown file not found at: {md_path}")
        return
        
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()
        
    doc = Document()
    
    # Set standard margins (1 inch)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    # Font Settings
    font_name = "TH Sarabun New"
    
    # Color definition
    primary_color = RGBColor(0, 86, 145)     # Deep Blue (#005691)
    secondary_color = RGBColor(0, 128, 128)   # Teal (#008080)
    dark_gray = RGBColor(80, 80, 80)
    text_color = RGBColor(33, 33, 33)
    
    # Split by lines
    lines = content.split('\n')
    
    in_table = False
    table_headers = []
    table_rows = []
    
    in_quote = False
    quote_type = None  # TIP, IMPORTANT
    quote_lines = []
    
    in_code = False
    code_lines = []
    code_lang = ""
    
    i = 0
    while i < len(lines):
        line = lines[i]
        line_stripped = line.strip()
        
        # Parse Code Blocks
        if line_stripped.startswith('```'):
            if not in_code:
                in_code = True
                code_lang = line_stripped[3:].strip()
                code_lines = []
                i += 1
                continue
            else:
                # End of code block, render it!
                tbl = doc.add_table(rows=1, cols=1)
                tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                tbl.autofit = False
                tbl.columns[0].width = Inches(6.5)
                
                cell = tbl.cell(0, 0)
                set_cell_background(cell, "F5F5F5") # Light gray
                set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
                set_cell_borders(cell, left=("888888", 12, 'single'))
                
                # Join code lines with newline
                code_content = "\n".join(code_lines)
                
                p = cell.paragraphs[0]
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.0
                
                run = p.add_run(code_content)
                run.font.name = "Consolas"
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(40, 40, 40)
                run._r.get_or_add_rPr().get_or_add_rFonts().set(qn('w:ascii'), "Consolas")
                run._r.get_or_add_rPr().get_or_add_rFonts().set(qn('w:hAnsi'), "Consolas")
                
                # Space after code block table
                doc.add_paragraph().paragraph_format.space_after = Pt(6)
                
                in_code = False
                code_lines = []
                i += 1
                continue
                
        if in_code:
            code_lines.append(line) # Keep original line formatting/indentation
            i += 1
            continue
            
        # Parse Callouts
        if line_stripped.startswith('>'):
            in_quote = True
            clean_line = re.sub(r'^>\s*', '', line_stripped)
            if '[!TIP]' in clean_line:
                quote_type = 'TIP'
                clean_line = clean_line.replace('[!TIP]', '').strip()
            elif '[!IMPORTANT]' in clean_line:
                quote_type = 'IMPORTANT'
                clean_line = clean_line.replace('[!IMPORTANT]', '').strip()
            
            if clean_line:
                quote_lines.append(clean_line)
            i += 1
            continue
        elif in_quote and not line_stripped.startswith('>'):
            # End of callout box, render it!
            bg_color = "E8F5E9" if quote_type == 'TIP' else "FFEBEE"
            border_color = "2E7D32" if quote_type == 'TIP' else "C62828"
            title_text = "💡 เคล็ดลับเพิ่มเติม (Tip):" if quote_type == 'TIP' else "⚠️ ข้อควรระวัง (Important):"
            title_color = RGBColor(46, 125, 50) if quote_type == 'TIP' else RGBColor(198, 40, 40)
            
            # Create a 1x1 table for the callout box
            tbl = doc.add_table(rows=1, cols=1)
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            tbl.autofit = False
            tbl.columns[0].width = Inches(6.5)
            
            cell = tbl.cell(0, 0)
            set_cell_background(cell, bg_color)
            set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
            set_cell_borders(cell, left=(border_color, 24, 'single'))
            
            # Add Title inside callout
            add_paragraph_with_runs(cell, title_text, size_pt=12.5, bold=True, color=title_color, space_after=4)
            
            # Add content inside callout
            for q_line in quote_lines:
                add_paragraph_with_runs(cell, q_line, size_pt=11, color=dark_gray, space_after=3)
                
            doc.add_paragraph().paragraph_format.space_after = Pt(6)
            
            in_quote = False
            quote_type = None
            quote_lines = []
            
        # Parse Table
        if line_stripped.startswith('|'):
            if not in_table:
                in_table = True
                table_headers = [c.strip() for c in line_stripped.split('|')[1:-1]]
                table_rows = []
                # Skip the separator line | :--- | :--- |
                i += 2
                continue
            else:
                row_cols = [c.strip() for c in line_stripped.split('|')[1:-1]]
                table_rows.append(row_cols)
                i += 1
                continue
        elif in_table and not line_stripped.startswith('|'):
            # End of table, render it!
            num_rows = len(table_rows) + 1
            num_cols = len(table_headers)
            
            tbl = doc.add_table(rows=num_rows, cols=num_cols)
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            tbl.autofit = False
            
            # Customize column widths based on table schema
            if num_cols == 4:
                col_widths = [Inches(1.5), Inches(1.5), Inches(1.8), Inches(1.7)]
            elif num_cols == 3:
                col_widths = [Inches(1.5), Inches(2.0), Inches(3.0)]
            else:
                col_widths = [Inches(6.5 / num_cols)] * num_cols
                
            for col_idx, width in enumerate(col_widths):
                if col_idx < len(tbl.columns):
                    tbl.columns[col_idx].width = width
                
            # Render Headers
            hdr_cells = tbl.rows[0].cells
            for col_idx, text in enumerate(table_headers):
                cell = hdr_cells[col_idx]
                cell.width = col_widths[col_idx]
                set_cell_background(cell, "005691") # Deep Blue
                set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
                set_cell_borders(cell, bottom=("002D54", 12, 'single'))
                
                p = cell.paragraphs[0]
                p.paragraph_format.space_after = Pt(2)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(text)
                run.font.name = font_name
                run.font.size = Pt(11)
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run._r.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), font_name)
                
            # Render Rows
            for row_idx, row_data in enumerate(table_rows):
                row_cells = tbl.rows[row_idx + 1].cells
                row_bg = "F4F8FA" if row_idx % 2 == 1 else "FFFFFF"
                
                for col_idx, text in enumerate(row_data):
                    if col_idx >= len(row_cells):
                        break
                    cell = row_cells[col_idx]
                    cell.width = col_widths[col_idx]
                    set_cell_background(cell, row_bg)
                    set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
                    set_cell_borders(cell, bottom=("E0E0E0", 4, 'single'))
                    
                    cell_text = text.replace("<br>", "\n").replace("<br/>", "\n")
                    text_parts = cell_text.split('\n')
                    
                    p = cell.paragraphs[0]
                    p.paragraph_format.space_after = Pt(4)
                    
                    for idx, part in enumerate(text_parts):
                        part = part.strip()
                        if not part:
                            continue
                        
                        is_bullet = False
                        if part.startswith('- '):
                            part = part[2:]
                            is_bullet = True
                            
                        if idx > 0 or is_bullet:
                            p = cell.add_paragraph()
                            p.paragraph_format.space_after = Pt(4)
                            
                        if is_bullet:
                            run_bullet = p.add_run("• ")
                            run_bullet.font.name = font_name
                            run_bullet.bold = True
                            run_bullet._r.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), font_name)
                            
                        parts = re.split(r'(\*\*.*?\*\*)', part)
                        for pt in parts:
                            if not pt:
                                continue
                            run = p.add_run()
                            run.font.name = font_name
                            run.font.size = Pt(10.5)
                            run._r.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), font_name)
                            run.font.color.rgb = text_color
                            
                            if pt.startswith('**') and pt.endswith('**'):
                                run.text = pt[2:-2]
                                run.bold = True
                            else:
                                run.text = pt
            
            doc.add_paragraph().paragraph_format.space_after = Pt(12)
            
            in_table = False
            table_headers = []
            table_rows = []
            
        # Parse Headings and Normal lines
        if not in_table and not in_quote and not in_code:
            if line_stripped.startswith('# '):
                title = line_stripped[2:]
                # Add Page break before Chapter titles to look very clean
                if doc.paragraphs:
                    doc.add_page_break()
                add_paragraph_with_runs(doc, title, size_pt=18, bold=True, color=primary_color, space_after=12)
            elif line_stripped.startswith('## '):
                h2 = line_stripped[3:]
                add_paragraph_with_runs(doc, h2, size_pt=14, bold=True, color=secondary_color, space_after=8)
            elif line_stripped.startswith('### '):
                h3 = line_stripped[4:]
                add_paragraph_with_runs(doc, h3, size_pt=12, bold=True, color=dark_gray, space_after=6)
            elif line_stripped.startswith('#### '):
                h4 = line_stripped[5:]
                add_paragraph_with_runs(doc, h4, size_pt=11.5, bold=True, color=dark_gray, space_after=4)
            elif line_stripped.startswith('---') and not line_stripped.startswith('----'):
                # Line separator
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                p_border = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="6" w:space="1" w:color="CCCCCC"/></w:pBdr>')
                p._p.get_or_add_pPr().append(p_border)
            elif line_stripped.startswith('* ') or line_stripped.startswith('- '):
                bullet_text = line_stripped[2:]
                p = doc.add_paragraph(style='List Bullet')
                p.paragraph_format.space_after = Pt(4)
                
                parts = re.split(r'(\*\*.*?\*\*)', bullet_text)
                for part in parts:
                    if not part:
                        continue
                    run = p.add_run()
                    run.font.name = font_name
                    run.font.size = Pt(11)
                    run._r.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), font_name)
                    run.font.color.rgb = text_color
                    
                    if part.startswith('**') and part.endswith('**'):
                        run.text = part[2:-2]
                        run.bold = True
                    else:
                        run.text = part
            elif line_stripped:
                # Normal paragraph
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.line_spacing = 1.15
                
                parts = re.split(r'(\*\*.*?\*\*)', line_stripped)
                for part in parts:
                    if not part:
                        continue
                    run = p.add_run()
                    run.font.name = font_name
                    run.font.size = Pt(11)
                    run._r.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), font_name)
                    run.font.color.rgb = text_color
                    
                    if part.startswith('**') and part.endswith('**'):
                        run.text = part[2:-2]
                        run.bold = True
                    else:
                        run.text = part
                        
        i += 1
        
    print(f"Saving beautifully formatted Word document to: {docx_path}")
    doc.save(docx_path)
    print("Done handbook conversion successfully!")

if __name__ == '__main__':
    main()
